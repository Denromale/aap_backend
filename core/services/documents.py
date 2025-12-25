from __future__ import annotations

import os
from io import BytesIO
from typing import Tuple

from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document

from core.models import Client, ClientDocument, ProcedureFile


def fill_docx(template_path: str, context: dict) -> BytesIO:
    """
    Открывает docx, подставляет значения по меткам и возвращает файл в памяти.
    Логика 1-в-1 как в views.py.
    """
    doc = Document(template_path)

    def replace_in_paragraph(paragraph):
        original_text = paragraph.text
        for key, value in context.items():
            if key in original_text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)
                if key in paragraph.text and original_text.strip() == key:
                    paragraph.text = original_text.replace(key, value)

    # --- Абзацы ---
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # --- Таблицы ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fill_docx_bytes(template_name: str, client: Client, user):
    """
    Генерирует DOCX из шаблона core/docs/<template_name>,
    подставляя плейсхолдеры вида {{ PLACEHOLDER }}.
    """
    template_path = os.path.join(settings.BASE_DIR, "core", "docs", template_name)

    def fmt_date(d):
        return d.strftime("%d.%m.%Y") if d else ""

    def fmt_decimal(x):
        if x is None:
            return ""
        try:
            return f"{x:,.2f}".replace(",", " ").replace(".", ",")
        except Exception:
            return str(x)

    def yn(v: bool):
        return "Так" if v else "Ні"

    # собрать полный адрес одной строкой
    address_parts = [
        client.address_zip,
        client.address_country,
        client.address_city,
        client.address_street,
        client.address_building,
        (f"оф. {client.address_office}" if client.address_office else None),
    ]
    address_full = ", ".join([p for p in address_parts if p])

    today_str = timezone.now().strftime("%d.%m.%Y")

    context_doc = {
        # базовое
        "{{ CLIENT_NAME }}": client.name or "",
        "{{ CLIENT_EDRPOU }}": client.edrpou or "",
        "{{ REPORTING_PERIOD }}": client.reporting_period or "",
        "{{ TODAY_DATE }}": today_str,
        "{{ CURRENT_USER }}": user.get_full_name() or user.username,

        # адрес
        "{{ CLIENT_ADDRESS_FULL }}": address_full,
        "{{ CLIENT_COUNTRY }}": client.address_country or "",
        "{{ CLIENT_CITY }}": client.address_city or "",
        "{{ CLIENT_STREET }}": client.address_street or "",
        "{{ CLIENT_BUILDING }}": client.address_building or "",
        "{{ CLIENT_OFFICE }}": client.address_office or "",
        "{{ CLIENT_ZIP }}": client.address_zip or "",

        # деятельность
        "{{ CLIENT_KVED }}": client.kved or "",
        "{{ MANDATORY_AUDIT }}": yn(bool(client.mandatory_audit)),

        # реквизиты договора/документа (если заполнены на карточке клиента)
        "{{ CONTRACT_NUMBER }}": client.requisites_number or "",
        "{{ CONTRACT_DATE }}": fmt_date(client.requisites_date),
        "{{ CONTRACT_AMOUNT }}": fmt_decimal(client.requisites_amount),
        "{{ CONTRACT_VAT }}": fmt_decimal(client.requisites_vat),
        "{{ CONTRACT_DEADLINE }}": fmt_date(client.contract_deadline),

        # предмет / детали задания
        "{{ ENGAGEMENT_SUBJECT }}": (client.get_engagement_subject_display() if client.engagement_subject else ""),
        "{{ PLANNED_HOURS }}": fmt_decimal(client.planned_hours),

        # уполномоченное лицо
        "{{ AUTH_PERSON_NAME }}": client.authorized_person_name or "",
        "{{ AUTH_PERSON_EMAIL }}": client.authorized_person_email or "",

        # команда
        "{{ MANAGER }}": client.manager.get_full_name() if client.manager else "",
        "{{ QA_MANAGER }}": client.qa_manager.get_full_name() if client.qa_manager else "",
        "{{ AUDITOR_1 }}": client.auditor.get_full_name() if client.auditor else "",
        "{{ AUDITOR_2 }}": client.auditor2.get_full_name() if client.auditor2 else "",
        "{{ AUDITOR_3 }}": client.auditor3.get_full_name() if client.auditor3 else "",
        "{{ ASSISTANT_1 }}": client.assistant.get_full_name() if client.assistant else "",
        "{{ ASSISTANT_2 }}": client.assistant2.get_full_name() if client.assistant2 else "",
        "{{ ASSISTANT_3 }}": client.assistant3.get_full_name() if client.assistant3 else "",
        "{{ ASSISTANT_4 }}": client.assistant4.get_full_name() if client.assistant4 else "",

        # данные по аудиторскому отчету (если используешь)
        "{{ AUDIT_REPORT_NUMBER }}": client.audit_report_number or "",
        "{{ AUDIT_REPORT_DATE }}": fmt_date(client.audit_report_date),
        "{{ AUDIT_REPORT_TYPE }}": (client.get_audit_report_type_display() if getattr(client, "audit_report_type", None) else ""),
        "{{ AUDIT_REPORT_PARAGRAPH }}": client.audit_report_paragraph or "",
    }

    file_obj = fill_docx(template_path, context_doc)
    return file_obj.getvalue()



def _step15_save_generated(
    request, *, client: Client, substep, file_bytes: bytes, filename: str, title: str
) -> Tuple[ProcedureFile, ClientDocument]:
    """
    Сохраняем один раз в storage через ProcedureFile,
    а ClientDocument ссылаем на тот же file.name.
    Логика 1-в-1 как в views.py.
    """
    # 1) ProcedureFile (для отображения на шаге)
    pf = ProcedureFile.objects.create(
        client=client,
        procedure_code=str(substep.id),
        uploaded_by=request.user,
        title=title or filename,
    )
    pf.file.save(filename, ContentFile(file_bytes), save=True)

    # 2) ClientDocument (База документів)
    step_label = f"Step {substep.step.order}.{substep.order}"

    doc = ClientDocument.objects.create(
        organization=request.organization,
        client=client,
        original_name=filename,
        doc_type="request",
        custom_label=step_label,
        uploaded_by=request.user,
    )

    # НЕ копируем файл, а ссылаемся на уже сохранённый pf.file
    doc.file.name = pf.file.name
    doc.save(update_fields=["file"])

    return pf, doc
